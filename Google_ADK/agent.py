"""
Google ADK agent with:
  - LangChain tool         (Wikipedia search via LangchainTool)
  - LangGraph workflow     (summarize → key points → sentiment analysis)
  - File-system tools      (read_file, list_directory, get_file_info, search_in_files)

Required packages:
    pip install python-docx pypdf openpyxl \
                langchain langchain-community langchain-google-genai \
                langgraph wikipedia
"""

import base64
import csv
import datetime
import io
import os
from pathlib import Path
from typing import TypedDict

from google.adk.agents.callback_context import CallbackContext
from google.adk.agents.llm_agent import Agent
from google.adk.models.llm_request import LlmRequest


# ── LangGraph workflow ────────────────────────────────────────────────────────
# A three-step sequential workflow: summarize → key points → sentiment.
# Each node calls Gemini via LangChain's ChatGoogleGenerativeAI.

from langgraph.graph import END, StateGraph
from langchain_google_genai import ChatGoogleGenerativeAI


class _AnalysisState(TypedDict):
    content: str
    summary: str
    key_points: str
    sentiment: str


_llm = ChatGoogleGenerativeAI(
    model="gemini-2.5-flash",
    google_api_key=os.getenv("GOOGLE_API_KEY"),
)


def _summarize_node(state: _AnalysisState) -> _AnalysisState:
    res = _llm.invoke(f"Summarize the following text in 2-3 sentences:\n\n{state['content']}")
    return {"summary": res.content}


def _key_points_node(state: _AnalysisState) -> _AnalysisState:
    res = _llm.invoke(
        f"Extract 3-5 key points as a numbered list from the following text:\n\n{state['content']}"
    )
    return {"key_points": res.content}


def _sentiment_node(state: _AnalysisState) -> _AnalysisState:
    res = _llm.invoke(
        f"What is the overall sentiment (positive / negative / neutral) of the following text and why?\n\n{state['content']}"
    )
    return {"sentiment": res.content}


_graph_builder = StateGraph(_AnalysisState)
_graph_builder.add_node("summarize", _summarize_node)
_graph_builder.add_node("extract_key_points", _key_points_node)
_graph_builder.add_node("analyze_sentiment", _sentiment_node)
_graph_builder.set_entry_point("summarize")
_graph_builder.add_edge("summarize", "extract_key_points")
_graph_builder.add_edge("extract_key_points", "analyze_sentiment")
_graph_builder.add_edge("analyze_sentiment", END)
_analysis_graph = _graph_builder.compile()


# ── LangChain tool ────────────────────────────────────────────────────────────
# Wraps a LangChain tool so ADK can call it like any other tool.

from google.adk.tools.langchain_tool import LangchainTool
from langchain_community.tools import WikipediaQueryRun
from langchain_community.utilities import WikipediaAPIWrapper

wikipedia_search = LangchainTool(
    tool=WikipediaQueryRun(api_wrapper=WikipediaAPIWrapper()),
)


# ── Inline-blob converters (used by before_model_callback) ───────────────────

def _docx_from_bytes(data: bytes) -> str:
    try:
        from docx import Document
    except ImportError:
        raise ImportError("Run: pip install python-docx")
    doc = Document(io.BytesIO(data))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            parts.append("\t".join(cell.text.strip() for cell in row.cells))
    return "\n".join(parts)


def _xlsx_from_bytes(data: bytes) -> str:
    try:
        import openpyxl
    except ImportError:
        raise ImportError("Run: pip install openpyxl")
    wb = openpyxl.load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    parts = []
    for sheet in wb.worksheets:
        parts.append(f"[Sheet: {sheet.title}]")
        for row in sheet.iter_rows(values_only=True):
            parts.append("\t".join("" if v is None else str(v) for v in row))
    return "\n".join(parts)


def _pdf_from_bytes(data: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError:
        raise ImportError("Run: pip install pypdf")
    reader = PdfReader(io.BytesIO(data))
    return "\n".join(page.extract_text() or "" for page in reader.pages)


_MIME_CONVERTERS = {
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": _docx_from_bytes,
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet": _xlsx_from_bytes,
    "application/vnd.ms-excel": _xlsx_from_bytes,
}


def _convert_unsupported_files(
    callback_context: CallbackContext,
    llm_request: LlmRequest,
) -> None:
    """Before-model callback: replace unsupported file blobs with extracted text."""
    from google.genai import types as genai_types

    if not llm_request.contents:
        return None

    for content in llm_request.contents:
        if not content.parts:
            continue
        new_parts = []
        for part in content.parts:
            inline = getattr(part, "inline_data", None)
            if inline and inline.mime_type in _MIME_CONVERTERS:
                raw = inline.data
                if isinstance(raw, str):
                    raw = base64.b64decode(raw)
                try:
                    text = _MIME_CONVERTERS[inline.mime_type](raw)
                    new_parts.append(
                        genai_types.Part(
                            text=f"[Extracted from uploaded {inline.mime_type} file]\n\n{text}"
                        )
                    )
                except Exception as exc:
                    new_parts.append(
                        genai_types.Part(text=f"[Could not extract file ({inline.mime_type}): {exc}]")
                    )
            else:
                new_parts.append(part)
        content.parts = new_parts

    return None


# ── File-system tools ─────────────────────────────────────────────────────────

def read_file(file_path: str) -> dict:
    """Read the contents of a file. Supports .txt, .py, .docx, .pdf, .xlsx, .csv, and more.

    Args:
        file_path: Absolute or relative path to the file.

    Returns:
        A dict with 'content', 'lines', and 'size_bytes', or 'error' on failure.
    """
    try:
        path = Path(file_path)
        if not path.exists():
            return {"error": f"File not found: {file_path}"}
        if not path.is_file():
            return {"error": f"Not a file: {file_path}"}

        ext = path.suffix.lower()
        data = path.read_bytes()

        if ext == ".docx":
            content = _docx_from_bytes(data)
        elif ext in (".xlsx", ".xls"):
            content = _xlsx_from_bytes(data)
        elif ext == ".pdf":
            content = _pdf_from_bytes(data)
        elif ext == ".csv":
            raw = data.decode("utf-8", errors="replace")
            content = "\n".join(", ".join(row) for row in csv.reader(io.StringIO(raw)))
        else:
            content = data.decode("utf-8", errors="replace")

        return {
            "file_path": str(path.resolve()),
            "file_type": ext or "unknown",
            "content": content,
            "lines": len(content.splitlines()),
            "size_bytes": path.stat().st_size,
        }
    except Exception as exc:
        return {"error": str(exc)}


def list_directory(directory_path: str) -> dict:
    """List files and subdirectories inside a directory.

    Args:
        directory_path: Path to the directory.

    Returns:
        A dict with 'files' and 'directories' lists, or 'error' on failure.
    """
    try:
        path = Path(directory_path)
        if not path.exists():
            return {"error": f"Directory not found: {directory_path}"}
        if not path.is_dir():
            return {"error": f"Not a directory: {directory_path}"}

        files, directories = [], []
        for item in sorted(path.iterdir()):
            if item.is_file():
                files.append({"name": item.name, "extension": item.suffix, "size_bytes": item.stat().st_size})
            elif item.is_dir():
                directories.append({"name": item.name})

        return {
            "directory": str(path.resolve()),
            "files": files,
            "directories": directories,
            "total_files": len(files),
            "total_directories": len(directories),
        }
    except Exception as exc:
        return {"error": str(exc)}


def get_file_info(file_path: str) -> dict:
    """Get metadata about a file or directory (size, timestamps, type).

    Args:
        file_path: Path to the file or directory.

    Returns:
        A dict with metadata, or 'error' on failure.
    """
    try:
        path = Path(file_path)
        if not path.exists():
            return {"error": f"Path not found: {file_path}"}
        stat = path.stat()
        return {
            "name": path.name,
            "extension": path.suffix,
            "absolute_path": str(path.resolve()),
            "is_file": path.is_file(),
            "is_directory": path.is_dir(),
            "size_bytes": stat.st_size,
            "size_kb": round(stat.st_size / 1024, 2),
            "created": datetime.datetime.fromtimestamp(stat.st_ctime).isoformat(),
            "modified": datetime.datetime.fromtimestamp(stat.st_mtime).isoformat(),
        }
    except Exception as exc:
        return {"error": str(exc)}


def search_in_files(directory: str, search_term: str, file_pattern: str = "*") -> dict:
    """Search for a text term across files in a directory (recursive, case-insensitive).

    Args:
        directory: Directory to search in.
        search_term: Text to look for.
        file_pattern: Glob filter, e.g. '*.py' or '*.txt'. Defaults to '*'.

    Returns:
        A dict with matched files and line numbers, or 'error' on failure.
    """
    try:
        path = Path(directory)
        if not path.exists():
            return {"error": f"Directory not found: {directory}"}

        results = []
        for fp in sorted(path.rglob(file_pattern)):
            if not fp.is_file():
                continue
            try:
                lines = fp.read_text(encoding="utf-8", errors="replace").splitlines()
                hits = [
                    {"line_number": i + 1, "line": line.strip()}
                    for i, line in enumerate(lines)
                    if search_term.lower() in line.lower()
                ]
                if hits:
                    results.append({"file": str(fp), "matches": hits})
            except Exception:
                continue

        return {
            "search_term": search_term,
            "directory": str(path.resolve()),
            "file_pattern": file_pattern,
            "files_with_matches": len(results),
            "results": results,
        }
    except Exception as exc:
        return {"error": str(exc)}


# ── LangGraph tool (exposed to ADK) ──────────────────────────────────────────

def analyze_document_workflow(content: str) -> dict:
    """Run a multi-step LangGraph analysis workflow on text content.

    The workflow runs three steps in sequence:
      1. Summarize  — produces a 2-3 sentence summary
      2. Key points — extracts a numbered list of key points
      3. Sentiment  — determines overall tone and explains why

    Use this after reading a file to get a structured deep analysis.

    Args:
        content: The text content to analyze (plain text, already extracted).

    Returns:
        A dict with 'summary', 'key_points', and 'sentiment' fields, or 'error' on failure.
    """
    try:
        result = _analysis_graph.invoke({
            "content": content,
            "summary": "",
            "key_points": "",
            "sentiment": "",
        })
        return {
            "summary": result.get("summary", ""),
            "key_points": result.get("key_points", ""),
            "sentiment": result.get("sentiment", ""),
        }
    except Exception as exc:
        return {"error": str(exc)}


# ── Agent ─────────────────────────────────────────────────────────────────────

root_agent = Agent(
    model="gemini-2.5-flash",
    name="root_agent",
    description=(
        "A helpful agent that can analyze uploaded files, explore the file system, "
        "search Wikipedia, and run deep document analysis via a LangGraph workflow."
    ),
    instruction=(
        "You are a helpful assistant with the following capabilities:\n"
        "- Read files (Word, Excel, PDF, CSV, plain text) and explore directories\n"
        "- Search Wikipedia for factual information using the wikipedia_search tool\n"
        "- Run deep document analysis (summary, key points, sentiment) using the "
        "  analyze_document_workflow tool — first read the file, then pass its content here\n"
        "When a file is uploaded or a path is given, use your tools to analyze it thoroughly. "
        "Always provide clear, detailed responses."
    ),
    before_model_callback=_convert_unsupported_files,
    tools=[
        # File-system tools
        read_file,
        list_directory,
        get_file_info,
        search_in_files,
        # LangGraph workflow
        analyze_document_workflow,
        # LangChain tool
        wikipedia_search,
    ],
)
